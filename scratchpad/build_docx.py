# -*- coding: utf-8 -*-
"""Build NEXVORA_SYSTEM_GUIDE_AR.docx from a markdown-like source.

Key bidi handling:
- Every paragraph gets w:bidi=1 + RIGHT alignment (except code blocks).
- Each paragraph's text is tokenized into Arabic vs Latin/code runs.
  * Arabic runs: w:rtl=1, font Arial.
  * Latin runs: wrapped with U+2066 (LRI) ... U+2069 (PDI) to isolate as LTR
    chunks so Word's bidi engine renders them contiguously; font Segoe UI
    (or Consolas for code).
- Bold via **word**; inline code via `code` (rendered mono).
- Fenced code blocks (```...```) become LTR, monospace paragraphs with grey
  shading, no bidi.
- Tables set to RTL direction; cells reuse the tokenizer.

Also runs light content normalization:
- Latin commas between Arabic chars -> Arabic comma (،).
- Latin question mark after Arabic -> Arabic question mark (؟).
"""
import sys, os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

FONT_AR = 'Arial'
FONT_LATIN = 'Segoe UI'
FONT_MONO = 'Consolas'

ARABIC_RE = re.compile(r'[؀-ۿݐ-ݿࢠ-ࣿﭐ-﷿ﹰ-﻿]')
# a token is either: fenced-inline-code `..`, bold **..**, latin/code chunk, or arabic chunk
TOKEN_RE = re.compile(
    r'(`[^`]+`)'                                # 1: inline code
    r'|(\*\*[^*]+\*\*)'                         # 2: bold
    r'|([A-Za-z0-9_@#\./\-+=<>:{}\[\]()\'"]+(?:\s+[A-Za-z0-9_@#\./\-+=<>:{}\[\]()\'"]+)*)'  # 3: latin run (may contain internal spaces)
    r'|([^`*A-Za-z0-9_@#\./\-+=<>:{}\[\]()\'"]+)'  # 4: everything else (arabic + punctuation)
)

LRI = '⁦'
PDI = '⁩'

# ---------------------------------------------------------------------------
# Low-level xml helpers
# ---------------------------------------------------------------------------

def _oxml(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    return el


def set_paragraph_bidi(paragraph, bidi=True):
    pPr = paragraph._p.get_or_add_pPr()
    # remove existing
    for tag in ('w:bidi',):
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)
    if bidi:
        pPr.append(_oxml('w:bidi', **{'w:val': '1'}))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _apply_run_props(run, *, font, size_pt, bold=False, rtl=False, mono=False,
                     color=None, shade=None):
    r = run._r
    rPr = r.get_or_add_rPr()
    # rFonts: set ascii/hAnsi/cs consistently
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _oxml('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    # size
    for tag in ('w:sz', 'w:szCs'):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
    rPr.append(_oxml('w:sz', **{'w:val': str(size_pt * 2)}))
    rPr.append(_oxml('w:szCs', **{'w:val': str(size_pt * 2)}))
    # bold
    if bold:
        rPr.append(_oxml('w:b'))
        rPr.append(_oxml('w:bCs'))
    # rtl
    if rtl:
        rPr.append(_oxml('w:rtl', **{'w:val': '1'}))
    # color
    if color is not None:
        col = _oxml('w:color', **{'w:val': '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])})
        rPr.append(col)
    # shading (for inline code)
    if shade is not None:
        shd = _oxml('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': shade})
        rPr.append(shd)


# ---------------------------------------------------------------------------
# Tokenization / run emission
# ---------------------------------------------------------------------------

def _add_run_text(paragraph, text, *, size_pt=11, bold=False, rtl=False,
                  mono=False, color=None, shade=None):
    if not text:
        return
    run = paragraph.add_run(text)
    font = FONT_MONO if mono else (FONT_AR if rtl else FONT_LATIN)
    _apply_run_props(run, font=font, size_pt=size_pt, bold=bold, rtl=rtl,
                     mono=mono, color=color, shade=shade)


def emit_paragraph_runs(paragraph, text, *, size_pt=11, base_bold=False,
                        base_color=None):
    """Split text into runs (arabic vs latin vs inline code vs bold) and add."""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        code, bold_chunk, latin, other = m.group(1), m.group(2), m.group(3), m.group(4)
        if code:
            inner = code[1:-1]
            _add_run_text(paragraph, LRI + inner + PDI, size_pt=size_pt,
                          bold=base_bold, rtl=False, mono=True,
                          shade='F2F2F2')
        elif bold_chunk:
            inner = bold_chunk[2:-2]
            # bold chunks can themselves be arabic or latin
            if ARABIC_RE.search(inner):
                _add_run_text(paragraph, inner, size_pt=size_pt, bold=True,
                              rtl=True, color=base_color)
            else:
                _add_run_text(paragraph, LRI + inner + PDI, size_pt=size_pt,
                              bold=True, rtl=False, color=base_color)
        elif latin:
            _add_run_text(paragraph, LRI + latin + PDI, size_pt=size_pt,
                          bold=base_bold, rtl=False, color=base_color)
        elif other:
            # other = arabic + shared punctuation
            _add_run_text(paragraph, other, size_pt=size_pt, bold=base_bold,
                          rtl=True, color=base_color)


# ---------------------------------------------------------------------------
# Content normalization
# ---------------------------------------------------------------------------

def normalize_arabic_punct(text):
    # Replace , between Arabic-adjacent contexts with Arabic comma
    def replace_comma(m):
        return '،'
    # Arabic , Arabic  or  Arabic , space
    text = re.sub(r'(?<=[؀-ۿ])\s*,\s*', '، ', text)
    # Arabic ? at end -> Arabic question mark
    text = re.sub(r'(?<=[؀-ۿ])\s*\?', ' ؟', text)
    # Semicolon between arabic
    text = re.sub(r'(?<=[؀-ۿ])\s*;\s*', '؛ ', text)
    return text


# ---------------------------------------------------------------------------
# Block builders
# ---------------------------------------------------------------------------

HEADING_COLORS = {
    1: RGBColor(0x1F, 0x3A, 0x5F),
    2: RGBColor(0x2E, 0x59, 0x8A),
    3: RGBColor(0x4A, 0x4A, 0x4A),
}
HEADING_SIZES = {1: 22, 2: 16, 3: 13}


def add_heading(doc, text, level=1):
    text = normalize_arabic_punct(text)
    p = doc.add_paragraph()
    set_paragraph_bidi(p, True)
    size = HEADING_SIZES.get(level, 11)
    color = HEADING_COLORS.get(level)
    emit_paragraph_runs(p, text, size_pt=size, base_bold=True,
                        base_color=(color[0], color[1], color[2]) if color else None)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    # Give H1 a subtle underline via bottom border
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = _oxml('w:pBdr')
        bottom = _oxml('w:bottom', **{'w:val': 'single', 'w:sz': '8', 'w:space': '1', 'w:color': '1F3A5F'})
        pBdr.append(bottom)
        pPr.append(pBdr)


def add_paragraph(doc, text, *, bullet=False):
    text = normalize_arabic_punct(text)
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
    set_paragraph_bidi(p, True)
    emit_paragraph_runs(p, text, size_pt=11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4


def add_code_block(doc, code):
    p = doc.add_paragraph()
    set_paragraph_bidi(p, False)
    # LTR alignment
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shd = _oxml('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': 'F5F5F5'})
    pPr.append(shd)
    # border
    pBdr = _oxml('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        pBdr.append(_oxml('w:' + side, **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': 'CCCCCC'}))
    pPr.append(pBdr)
    _add_run_text(p, code, size_pt=9, rtl=False, mono=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    # pad rows
    rows = [r + [''] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    try:
        table.table_direction = WD_TABLE_DIRECTION.RTL
    except Exception:
        pass
    # also set bidiVisual on tblPr for safety
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr.find(qn('w:bidiVisual')) is None:
        tblPr.append(_oxml('w:bidiVisual', **{'w:val': '1'}))
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph_bidi(p, True)
            text = normalize_arabic_punct(cell_text)
            emit_paragraph_runs(p, text, size_pt=10, base_bold=(i == 0))
            # header shading
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = _oxml('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': 'E8ECF1'})
                tcPr.append(shd)


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

def parse_and_write(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        # fenced code block
        if line.strip().startswith('```'):
            j = i + 1
            code_lines = []
            while j < len(lines) and not lines[j].strip().startswith('```'):
                code_lines.append(lines[j])
                j += 1
            add_code_block(doc, '\n'.join(code_lines))
            i = j + 1
            continue
        # table
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                parts = [c.strip() for c in tl.strip().strip('|').split('|')]
                if all(re.match(r'^:?-+:?$', p) for p in parts if p):
                    continue
                rows.append(parts)
            add_table(doc, rows)
            continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith('- ') or line.startswith('* '):
            add_paragraph(doc, line[2:].strip(), bullet=True)
        elif line.strip() == '':
            pass
        else:
            add_paragraph(doc, line.strip())
        i += 1


# ---------------------------------------------------------------------------
# Cover, footer
# ---------------------------------------------------------------------------

def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        set_paragraph_bidi(p, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = p.add_run()
        _apply_run_props(run, font=FONT_LATIN, size_pt=9, rtl=False)
        fldChar1 = _oxml('w:fldChar', **{'w:fldCharType': 'begin'})
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fldChar2 = _oxml('w:fldChar', **{'w:fldCharType': 'end'})
        run._r.append(fldChar1)
        run._r.append(instr)
        run._r.append(fldChar2)


def add_cover(doc):
    # top spacing
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    set_paragraph_bidi(title, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_text(title, 'دليل نظام NEXVORA الشامل', size_pt=34, bold=True,
                  rtl=True, color=(0x1F, 0x3A, 0x5F))
    sub = doc.add_paragraph()
    set_paragraph_bidi(sub, True)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_text(sub, 'دليل عملي شامل للمؤسس والشريك — من الصفر إلى الاحتراف',
                  size_pt=14, rtl=True, color=(0x4A, 0x4A, 0x4A))
    date_p = doc.add_paragraph()
    set_paragraph_bidi(date_p, True)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_text(date_p, '٢٠٢٦', size_pt=12, rtl=True,
                  color=(0x4A, 0x4A, 0x4A))
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    src = sys.argv[1]
    dst = sys.argv[2]
    with open(src, 'r', encoding='utf-8') as f:
        md = f.read()

    doc = Document()
    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_AR
    style.font.size = Pt(11)
    # Default paragraph format
    pf = style.paragraph_format
    pf.line_spacing = 1.4

    # Section: set bidi + margins
    sect = doc.sections[0]
    sect.left_margin = Cm(2.2)
    sect.right_margin = Cm(2.2)
    sect.top_margin = Cm(2.2)
    sect.bottom_margin = Cm(2.2)
    sectPr = sect._sectPr
    if sectPr.find(qn('w:bidi')) is None:
        sectPr.append(_oxml('w:bidi'))
    # RTL gutter
    if sectPr.find(qn('w:rtlGutter')) is None:
        sectPr.append(_oxml('w:rtlGutter'))

    add_cover(doc)
    parse_and_write(doc, md)
    add_page_number_footer(doc)

    doc.save(dst)
    size = os.path.getsize(dst)
    print(f'Saved {dst} ({size} bytes)')


if __name__ == '__main__':
    main()
